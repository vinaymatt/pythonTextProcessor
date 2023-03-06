#Description: This file is used to generate embeddings for the abstracts using the BERT model. The embeddings are then used to identify common themes and topics across the two fields.
from sentence_transformers import SentenceTransformer, util
from collections import Counter
import torch
import pickle
import os
import docx
from docx.shared import Pt
from tqdm import tqdm
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

os.environ['KMP_DUPLICATE_LIB_OK']='True'
torch.set_num_threads(4)
if torch.cuda.is_available():
    device = torch.device("cuda")
else:
    device = torch.device("cpu")
print("Using device:", device)

with open('/Users/vinay/PycharmProjects/pythonTextProcessor/data/Biochemlarge/Abstracts_unprocessed.txt', 'r') as f:
    text_biochemone = f.read()
abstracts_biochemone = text_biochemone.split('\n\n')
abstracts_biochemone = [a.strip() for a in abstracts_biochemone if a.strip()]

with open('//Users/vinay/PycharmProjects/pythonTextProcessor/data/Biomechlarge/Abstracts_unprocessed.txt', 'r') as f:
    text_biomechone = f.read()
abstracts_biomechone = text_biomechone.split('\n\n')
abstracts_biomechone = [a.strip() for a in abstracts_biomechone if a.strip()]

# Load the embeddings from a file
with open('/Users/vinay/PycharmProjects/pythonTextProcessor/data/embeddingsword2vec_biochem.pkl', 'rb') as f:
    biochem_embeddings = pickle.load(f)

with open('/Users/vinay/PycharmProjects/pythonTextProcessor/data/embeddingsword2vec_biomech.pkl', 'rb') as f:
    biomech_embeddings = pickle.load(f)

# Compare each biochemistry abstract embedding to all biomechanics abstract embeddings
similarities = {}
for i, biochem_embedding in enumerate(biochem_embeddings):
    max_similarities = []
    for j, biomech_embedding in enumerate(biomech_embeddings):
        similarity = cosine_similarity(biochem_embedding.reshape(1,-1), biomech_embedding.reshape(1,-1))[0][0]
        max_similarities.append((j, similarity))
    max_similarities.sort(key=lambda x: x[1], reverse=True)
    similarities[i] = max_similarities[:10]


# Save the top 10 most similar biochemistry abstracts to a word document file
doc = docx.Document()
for i, abstract in enumerate(abstracts_biomechone):
    doc.add_heading(f'Biomechanics Abstract {i+1}', level=1)
    for j, similarity in similarities.items():
        if i in [x[0] for x in similarity]:
            doc.add_heading(f'Biochemistry Abstract {j+1} (Similarity Score: {similarities[j][[x[0] for x in similarity].index(i)][1]:.4f})', level=2)
            doc.add_paragraph(abstracts_biochemone[j])
    doc.add_page_break()
'''
# Identify common themes and topics across the two fields
top_k = 5
common_words = set()
for i in range(len(abstracts_biochem)):
    query_embedding = embeddings_biochem[i]
    cos_scores = util.pytorch_cos_sim(query_embedding, embeddings_biomech)[0]
    cos_scores = cos_scores.cpu()

    # Find the indices of the top-k most similar abstracts from the biomechanics set
    top_results = torch.topk(cos_scores, k=top_k)

    for score, idx in zip(top_results[0], top_results[1]):
        words = abstracts_biomech[idx].split()
        word_counts = Counter(words)

        # Identify the most frequent words in the similar biomechanics abstracts
        most_common_words = set([word for word, count in word_counts.most_common(10)])

        # Add the most common words to the set of common words
        if common_words:
            common_words = common_words.intersection(most_common_words)
        else:
            common_words = most_common_words

# Add the common themes and topics to the Word document
common_themes = document.add_paragraph()
common_themes.add_run("\n\nCommon themes and topics across the two fields:").bold = True
common_themes.add_run(f" {common_words}")

# Generate hypotheses for future research
top_k = 5
hypotheses = document.add_paragraph()
hypotheses.add_run("\n\nHypotheses for future research:").bold

'''

doc.save('/Users/vinay/PycharmProjects/pythonTextProcessor/data/Word2Vecabstract_embeddings.docx')
