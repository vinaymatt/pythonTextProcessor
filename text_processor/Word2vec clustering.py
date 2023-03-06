import pickle
import numpy as np
from sklearn.decomposition import PCA
from sklearn.neighbors import BallTree
from sklearn.neighbors import NearestNeighbors
from umap import UMAP
from docx import Document
from tqdm import tqdm
from annoy import AnnoyIndex

# Load the embeddings from the pickle file
with open('/Users/vinay/PycharmProjects/pythonTextProcessor/data/Biochemlarge/word2vec_withTFID_biochem_embeddings.pkl', 'rb') as f:
    biochem_embeddings = pickle.load(f)

with open('/Users/vinay/PycharmProjects/pythonTextProcessor/data/Biomechlarge/word2vec_withTFID_biomech_embeddings.pkl', 'rb') as f:
    biomech_embeddings = pickle.load(f)

# Load the original abstracts
with open('/Users/vinay/PycharmProjects/pythonTextProcessor/data/Biochemlarge/word2vec_withTFID_original_filtered_abstracts_biochem.txt', 'r') as f:
    filtered_biochem_original_abstracts = f.read().split('\n\n')

with open('/Users/vinay/PycharmProjects/pythonTextProcessor/data/Biomechlarge/word2vec_withTFID_original_filtered_abstracts_biomech.txt', 'r') as f:
    filtered_biomech_original_abstracts = f.read().split('\n\n')

# Combine the embeddings
combined_embeddings = np.concatenate((biochem_embeddings, biomech_embeddings), axis=0)

# Dimensionality reduction using UMAP
reduced_embeddings = UMAP(n_neighbors=30, n_components=50).fit_transform(combined_embeddings)

# Perform nearest neighbors search
biochem_count = len(biochem_embeddings)
biomech_count = len(biomech_embeddings)
embedding_dim = reduced_embeddings.shape[1]
annoy_index = AnnoyIndex(embedding_dim, 'euclidean')
for i, embedding in enumerate(reduced_embeddings[biochem_count:]):
    annoy_index.add_item(i, embedding)
annoy_index.build(50)

nearest_distances = []
nearest_indices = []
for i in tqdm(range(biochem_count)):
    indices, distances = annoy_index.get_nns_by_vector(reduced_embeddings[i], biomech_count, include_distances=True)
    nearest_distances.append(distances)
    nearest_indices.append(indices)


nearest_biomech_indices = np.array(nearest_indices)

# Find the most similar abstract pairs
similar_pairs = []
for i, neighbors in enumerate(tqdm(nearest_biomech_indices, desc="Finding similar pairs")):
    for j, neighbor in enumerate(neighbors):
        distance = nearest_distances[i][j]
        similar_pairs.append((i, neighbor, distance))

# Sort the pairs by distance
similar_pairs.sort(key=lambda x: x[2])

# Create a new Word document
doc = Document()

# Write the top 10 most similar pairs to the Word document
for i, (biochem_idx, biomech_idx, distance) in tqdm(enumerate(similar_pairs[:10])):
    doc.add_heading(f'Most Similar Pair {i + 1}', level=1)
    doc.add_heading('Biochemistry Abstract', level=2)
    doc.add_paragraph(filtered_biochem_original_abstracts[biochem_idx])
    doc.add_heading('Biomechanics Abstract', level=2)
    doc.add_paragraph(filtered_biomech_original_abstracts[biomech_idx])
    doc.add_heading('Distance', level=2)
    doc.add_paragraph(str(distance))

doc.save('/Users/vinay/PycharmProjects/pythonTextProcessor/data/most_similar_abstractscleanedwithTFID.docx')