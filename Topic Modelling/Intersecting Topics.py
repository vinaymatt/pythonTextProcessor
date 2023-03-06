import os
import numpy as np
from gensim.corpora import Dictionary
from gensim.models import HdpModel
from scipy.spatial.distance import jensenshannon
import docx
def load_hdp_models(output_directory, hdp_model_file, dict_file):
    hdp_model = HdpModel.load(os.path.join(output_directory, hdp_model_file))
    dictionary = Dictionary.load(os.path.join(output_directory, dict_file))
    hdp_model.id2word = dictionary
    return hdp_model

def calculate_similarity(hdp_model1, hdp_model2):
    topics1 = hdp_model1.get_topics()
    topics2 = hdp_model2.get_topics()

    # Reshape the arrays to match their shapes
    if topics1.shape[1] < topics2.shape[1]:
        padding = np.zeros((topics1.shape[0], topics2.shape[1] - topics1.shape[1]))
        topics1 = np.hstack((topics1, padding))
    elif topics1.shape[1] > topics2.shape[1]:
        padding = np.zeros((topics2.shape[0], topics1.shape[1] - topics2.shape[1]))
        topics2 = np.hstack((topics2, padding))

    similarity_matrix = np.zeros((len(topics1), len(topics2)))

    for i, topic1 in enumerate(topics1):
        for j, topic2 in enumerate(topics2):
            similarity_matrix[i, j] = 1 - jensenshannon(topic1, topic2)

    return similarity_matrix

def find_intersecting_topics(similarity_matrix, hdp_model1, hdp_model2, threshold=0.5, num_words=10):
    intersecting_topic_indices = np.argwhere(similarity_matrix > threshold)
    intersecting_topics = []
    for i, j in intersecting_topic_indices:
        topic1_word_probs = hdp_model1.show_topic(i, topn=num_words)
        topic1_words = [word for word, _ in topic1_word_probs]
        topic2_word_probs = hdp_model2.show_topic(j, topn=num_words)
        topic2_words = [word for word, _ in topic2_word_probs]
        intersecting_topics.append((i, j, topic1_words, topic2_words))
    return intersecting_topics


def write_to_word(intersecting_topics):
    # Create a new Word document
    doc = docx.Document()

    # Add a heading to the document
    doc.add_heading("Intersecting Topics", level=1)

    # Loop through the intersecting topics and add each one to the document
    for i, j, topic1_words, topic2_words in intersecting_topics:
        # Add a sub-heading for each intersecting topic
        doc.add_heading(f"({i}, {j})", level=2)

        # Add the words for each topic
        doc.add_paragraph("Topic 1 words:")
        for word in topic1_words:
            doc.add_paragraph(word, style='List Bullet')

        doc.add_paragraph("Topic 2 words:")
        for word in topic2_words:
            doc.add_paragraph(word, style='List Bullet')

        # Add a page break between each intersecting topic
        doc.add_page_break()

    # Save the document
    doc.save("/Users/vinay/PycharmProjects/pythonTextProcessor/data/intersecting_topics.docx")

def main():
    output_directory = "/Users/vinay/PycharmProjects/pythonTextProcessor/data"
    biomech_dict_file = "biomech_dictionary.dict"
    biochem_dict_file = "biochem_dictionary.dict"

    # Load trained HDP models and dictionaries
    biomech_hdp_model = load_hdp_models(output_directory, "biomech_hdp_model.hdp", biomech_dict_file)
    biochem_hdp_model = load_hdp_models(output_directory, "biochem_hdp_model.hdp", biochem_dict_file)

    # Calculate topic similarity
    similarity_matrix = calculate_similarity(biomech_hdp_model, biochem_hdp_model)

    # Find intersecting topics
    intersecting_topics = find_intersecting_topics(similarity_matrix, biomech_hdp_model, biochem_hdp_model)

    # Write the intersecting topics to a Word document
    write_to_word(intersecting_topics)

if __name__ == "__main__":
    main()
